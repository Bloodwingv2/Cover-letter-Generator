import os
import time
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import subprocess
import sys

def get_job_description_from_cli():
    """Gets the job description from a command-line argument or prompts for multi-line input."""
    if len(sys.argv) > 1:
        # Get job description from command-line argument
        return ' '.join(sys.argv[1:])
    else:
        # No CLI argument given – read from stdin (multi-line input)
        print("Paste the full job description below. Press Ctrl+D (or Ctrl+Z on Windows) when done:")
        return sys.stdin.read()

def generate_cover_letter(job_description, api_key):
    """Generates a cover letter using the Groq API and removes the introductory sentence."""
    client = Groq(api_key=api_key)
    completion = client.chat.completions.create(
        model="llama3-70b-8192",
        messages=[
            {
                "role": "user",
                "content": f"""Generate a professional cover letter based on this job description. The cover letter must be exactly 1 page, maintain a formal, enthusiastic, and polished tone, and preserve placeholders like [Organization Name], [Email Address], [Phone Number], and [Name].

Structure the letter into 3 main paragraphs:
1. Opening paragraph: Explain how the applicant found out about the job and express interest in applying.
2. Second paragraph: Introduce the applicant briefly, mention a relevant event or experience, and highlight key skills or attributes.
3. Final paragraph: Share a past experience (e.g., volunteering), describe technical and soft skills, and express confidence in their suitability for the role. Add a closing paragraph inviting the reader for a personal meeting and providing contact options.

The salutation should begin with 'Dear [Hiring Manager Name],'. The closing should be 'Best regards,' followed by '[Your Name]'.

Ensure to use placeholders like [Your Name], [Your Address], [Your City, Postal Code], [Your Email Address], [Your Phone Number], [Your LinkedIn Profile], [Date], [Hiring Manager Name], [Hiring Manager Title], [Company Name], [Company Address], [Company City, Postal Code], [Job Title] where appropriate, and do not fill them in.

Job Description:

{job_description}"""
            }
        ],
        temperature=0.7,
        max_tokens=1024,
        top_p=1,
        stream=False,
        response_format={"type": "text"},
        stop=None,
    )
    
    generated_text = completion.choices[0].message.content
    unwanted_sentence = "Here is a professional cover letter based on the job description:"
    
    if generated_text.strip().startswith(unwanted_sentence):
        return generated_text.strip()[len(unwanted_sentence):].strip()
    else:
        return generated_text

import time

def create_docx(cover_letter_text):
    """Creates a .docx file with professional formatting, ensuring it fits on one page."""
    document = Document()
    
    # Set margins (in inches)
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set default font
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set line spacing for the entire document
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph_format.space_after = Pt(0)

    # Add Applicant Header Box
    # Create a table for the header to act as a box
    table = document.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5) # Adjust width as needed

    # Add content to the cell
    cell = table.rows[0].cells[0]
    
    # Name
    name_paragraph = cell.add_paragraph()
    name_run = name_paragraph.add_run('[Your Name]')
    name_run.bold = True
    name_run.font.size = Pt(24) # Larger font for name
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    contact_info_paragraph = cell.add_paragraph()
    contact_info_run = contact_info_paragraph.add_run('[Your Address], [Your City, Postal Code]\n[Your Email Address] | [Your Phone Number] | [Your LinkedIn Profile]')
    contact_info_run.font.size = Pt(12) # Slightly larger font for contact info
    contact_info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add some space after the header box
    document.add_paragraph().add_run().add_break()

    # Add Date (handled by Groq API content)
    # Add Recipient Information (handled by Groq API content)

    # Add content with compact spacing
    # Assuming the Groq API will provide the salutation and closing within cover_letter_text
    # We will split the text and add paragraphs
    for line in cover_letter_text.split('\n'):
        paragraph = document.add_paragraph(line)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save the document with a unique name
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    output_path = f"output/cover_letter_{timestamp}.docx"
    document.save(output_path)
    return output_path

def convert_to_pdf(docx_path):
    """Converts a .docx file to .pdf using LibreOffice."""
    output_dir = "output"
    try:
        subprocess.run(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            check=True,
            capture_output=True,
            text=True
        )
        return os.path.join(output_dir, "cover_letter.pdf")
    except FileNotFoundError:
        print("Error: 'soffice' command not found. Please ensure LibreOffice is installed and in your PATH.")
        return None
    except subprocess.CalledProcessError as e:
        print(f"Error during PDF conversion: {e.stderr}")
        return None

def main():
    """Main function to generate the cover letter."""
    load_dotenv()
    api_key = os.getenv("GROQ_API_KEY")

    if not api_key:
        print("Error: GROQ_API_KEY not found in .env file.")
        return

    job_description = get_job_description_from_cli()
    cover_letter_text = generate_cover_letter(job_description, api_key)
    
    print("\nGenerating cover letter...")
    
    docx_path = create_docx(cover_letter_text)
    print(f"Successfully created DOCX: {docx_path}")
    
    pdf_path = convert_to_pdf(docx_path)
    if pdf_path:
        print(f"Successfully created PDF: {pdf_path}")

if __name__ == "__main__":
    main()